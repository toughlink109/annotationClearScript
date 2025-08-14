import os
import re
from docx import Document

INPUT_FOLDER = './input_docs'
OUTPUT_FOLDER = './output_docs'
# REPLACE_SINGLE = '[注释已删除]'
# REPLACE_MULTI = '[多行注释已删除]'
REPLACE_SINGLE = ''  # 删除单行注释，不做标记
REPLACE_MULTI = ''   # 删除多行注释，不做标记

def replace_single_line_comments(line):
    return re.sub(r'//.*', REPLACE_SINGLE, line)


def replace_multiline_block_comments(text):
    return re.sub(r'/\*\*.*?\*/', REPLACE_MULTI, text, flags=re.DOTALL)


def remove_empty_lines(lines):
    """
    删除空行（只含空格或换行的行）
    """
    return [line for line in lines if line.strip() != '']


def process_docx_file(filepath, output_path):
    doc = Document(filepath)
    all_lines = [para.text for para in doc.paragraphs]
    text_joined = "\n".join(all_lines)

    # 替换注释
    text_joined = replace_multiline_block_comments(text_joined)
    processed_lines = [replace_single_line_comments(line) for line in text_joined.split('\n')]

    # 删除空行
    processed_lines = remove_empty_lines(processed_lines)

    # 保存新文档
    # 清空旧段落
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # 添加新段落
    for line in processed_lines:
        doc.add_paragraph(line)

    doc.save(output_path)


def batch_process_docs(input_dir, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for filename in os.listdir(input_dir):
        if filename.endswith('.docx'):
            input_path = os.path.join(input_dir, filename)
            output_path = os.path.join(output_dir, f'处理后_{filename}')
            print(f"处理文件: {filename}")
            process_docx_file(input_path, output_path)

    print("所有文件已处理完毕")


if __name__ == "__main__":
    batch_process_docs(INPUT_FOLDER, OUTPUT_FOLDER)